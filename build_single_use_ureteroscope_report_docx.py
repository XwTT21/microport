from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DATA_JSON = ROOT / "outputs" / "single_use_ureteroscope_screening" / "single_use_ureteroscope_screening_data.json"
OUTPUT = ROOT / "outputs" / "single_use_ureteroscope_screening" / "一次性输尿管镜文献筛选与评价报告_完整版.docx"


def clean(text: str, limit: int | None = None) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    if limit and len(text) > limit:
        return text[: limit - 1].rstrip() + "..."
    return text


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold=False, color="111827", size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(clean(str(text)))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_geometry(table, widths: list[float] | None) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    if not widths:
        return
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])


def repeat_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    repeat_header_row(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="0B2545", size=font_size)
        set_cell_shading(hdr[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("111827")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Single-use / Disposable Ureteroscope 文献检索、筛选与评价报告")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"基于本地Cochrane、Google Scholar、PubMed及二次搜索文献 | 生成日期：{date.today().isoformat()}")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("4B5563")
    doc.add_paragraph()


def build_report() -> None:
    data = json.loads(DATA_JSON.read_text(encoding="utf-8"))
    summary = data["summary"]
    report = data["report_summary"]
    app_counts = summary.get("appraisal_overall_counts", {})
    dataset_counts = report.get("dataset_counts", {})
    decision_counts = summary.get("decision_counts", {})

    doc = Document()
    setup_styles(doc)
    add_title(doc)

    doc.add_heading("1. Executive Summary / 摘要", level=1)
    doc.add_paragraph(
        "本报告按照一次性/可抛弃输尿管镜主题口径，对本地Cochrane、Google Scholar、PubMed及二次搜索下载文献进行去重、筛选、数据集分层、适应性评估和安全性/性能证据摘录。"
    )
    doc.add_paragraph(
        "筛选过程中复用既有PDF文本和PubMed题录缓存，并采用保守题录匹配策略：低置信题录匹配时优先依据PDF文件名和PDF文本，不强行套用PubMed题录。"
    )
    add_table(
        doc,
        ["指标", "结果"],
        [
            ["缓存PDF总数", str(summary.get("cache_pdf_count", summary.get("pdf_count", 0)))],
            ["产品文献排除PDF数", str(summary.get("product_pdf_excluded_count", 0))],
            ["实际进入筛选PDF数", str(summary.get("pdf_count", 0))],
            ["去重后筛选记录", str(summary.get("screened_unique_count", 0))],
            ["纳入 / 排除 / 待确认", f"{decision_counts.get('纳入', 0)} / {decision_counts.get('排除', 0)} / {decision_counts.get('待人工确认', 0)}"],
            ["相似器械临床数据", str(dataset_counts.get("Similar device clinical dataset", 0))],
            ["SOTA数据", str(dataset_counts.get("SOTA dataset", 0))],
            ["技术支持数据", str(dataset_counts.get("Technical supportive data", 0))],
            ["经济/资源证据", str(dataset_counts.get("Economic/resource evidence", 0))],
            ["病例支持证据", str(dataset_counts.get("Supportive case evidence", 0))],
            ["待确认分层", str(dataset_counts.get("Pending confirmation", 0))],
        ],
        widths=[2.3, 4.2],
        font_size=9,
    )

    doc.add_heading("2. Literature Search Scope and Source / 检索范围和来源", level=1)
    add_bullet(doc, "纳入本地Cochrane、Google Scholar、PubMed、文献2次搜索(24-26.6)中的PDF和PubMed TXT题录缓存。")
    add_bullet(doc, "路径或检索文件夹含“产品”的记录在输入阶段排除，不进入筛选总表。")
    add_bullet(doc, "PubMed TXT题录用于补充题名、作者、期刊、年份、DOI、PMID和摘要；低置信匹配不强行合并。")
    add_bullet(doc, "筛选.docx和已完成UAS流程仅作为流程、表格和报告结构参考，主题标准已改写为一次性输尿管镜。")

    doc.add_heading("3. Screening Criteria and Process / 筛选标准和流程", level=1)
    doc.add_paragraph(
        "纳入条件：题名或摘要显示single-use ureteroscope、disposable ureteroscope、single-use/disposable flexible ureteroscope、single-use/disposable ureterorenoscope、reusable vs single-use ureteroscope、digital disposable ureteroscope或中文同义词为研究对象、干预、对照、器械、技术平台或主要临床问题，并报告安全性、性能、有效性、资源或经济学相关结局。"
    )
    doc.add_paragraph(
        "排除条件：仅研究UAS、激光、PCNL、ESWL、取石篮、支架、药物、影像预测模型等，且一次性输尿管镜不是核心主题；或一次性输尿管镜仅作为方法中顺带使用的器械而非研究问题。无法可靠判断者进入待人工确认。"
    )
    add_table(
        doc,
        ["流程步骤", "数量", "说明"],
        [[r.get("Step", ""), str(r.get("Count", "")), clean(r.get("Description", ""), 150)] for r in data.get("flow_rows", [])],
        widths=[1.8, 0.8, 3.9],
        font_size=8,
    )

    doc.add_heading("4. Dataset Assignment / SOTA与相似器械分层", level=1)
    doc.add_paragraph("本报告将一次性输尿管镜证据分为以下数据集，便于后续临床评价、SOTA论证、技术支持和经济资源分析。")
    add_table(
        doc,
        ["数据集", "数量", "用途"],
        [
            ["Similar device clinical dataset", str(dataset_counts.get("Similar device clinical dataset", 0)), "直接评价一次性输尿管镜或相似一次性软镜的临床安全性/性能。"],
            ["SOTA dataset", str(dataset_counts.get("SOTA dataset", 0)), "系统综述、Meta、指南、专家共识或叙述综述，用于技术现状和总体背景。"],
            ["Technical supportive data", str(dataset_counts.get("Technical supportive data", 0)), "体外、工程、成像、弯曲、灌注、可用性和性能测试等非临床数据。"],
            ["Economic/resource evidence", str(dataset_counts.get("Economic/resource evidence", 0)), "一次性/可重复使用成本、维修、灭菌、流程、资源利用或环境影响。"],
            ["Supportive case evidence", str(dataset_counts.get("Supportive case evidence", 0)), "病例或特殊场景补充证据。"],
            ["Pending confirmation", str(dataset_counts.get("Pending confirmation", 0)), "主题权重或证据用途需人工复核。"],
        ],
        widths=[2.2, 0.7, 3.6],
        font_size=8,
    )

    doc.add_heading("5. Suitability and Contribution Appraisal / 适应性和贡献评价", level=1)
    doc.add_paragraph(
        "适应性评估采用一次性输尿管镜改写版D/A/P/R/T/O/S/C框架，分别评价器械相关性、用途一致性、患者适用性、报告质量、证据类型、结局相关性、统计支持和临床意义。"
    )
    add_table(
        doc,
        ["综合评价", "数量", "解释"],
        [
            ["高/核心证据", str(app_counts.get("高/核心证据", 0)), "器械、用途、人群、报告和结局均与一次性输尿管镜临床评价高度匹配。"],
            ["中等/可作为支持证据", str(app_counts.get("中等/可作为支持证据", 0)), "主题相关，但研究类型、人群、报告质量或统计/临床意义存在限制。"],
            ["低/仅支持性", str(app_counts.get("低/仅支持性", 0)), "主题或可审计性存在明显限制，仅作背景或支持性资料。"],
            ["待确认", str(app_counts.get("待确认", 0)), "需人工复核后决定是否纳入正式证据包。"],
        ],
        widths=[1.7, 0.7, 4.1],
        font_size=8,
    )

    doc.add_heading("6. Safety and Performance Evidence Summary / 安全性和性能证据总结", level=1)
    doc.add_paragraph(
        "自动数据摘录显示，一次性输尿管镜相关文献主要围绕结石清除率、并发症、感染/发热/脓毒症/SIRS、手术或操作时间、镜体故障、弯曲、灌注、图像质量、可用性、维修/灭菌/污染、成本、资源利用、环境影响、住院和再入院等维度报告安全性、性能或资源结局。详细逐篇摘录见Excel工作簿“安全性能数据提取”sheet。"
    )
    top_evidence = [r for r in data.get("performance_safety_extraction", []) if r.get("Dataset_Assignment") == "Similar device clinical dataset"][:10]
    if top_evidence:
        add_table(
            doc,
            ["ID", "题名", "设计/样本", "主要自动摘录结局"],
            [
                [
                    r.get("Record_ID", ""),
                    clean(r.get("Title", ""), 110),
                    clean(" / ".join([r.get("Study_Design", ""), r.get("Sample_Size_N", "")]), 60),
                    clean(" | ".join([r.get("Performance_Outcomes", ""), r.get("Safety_Outcomes", ""), r.get("Efficiency_or_Resource_Outcomes", "")]), 180),
                ]
                for r in top_evidence
            ],
            widths=[0.55, 2.35, 1.15, 2.45],
            font_size=7,
        )
    doc.add_paragraph("注：安全性/性能数据为自动抽取，核心证据正式提交前建议全文核对。")

    doc.add_heading("7. Included Evidence List / 纳入文献列表", level=1)
    appraisal_by_id = {r["Record_ID"]: r for r in data.get("appraisal", [])}
    included_rows = []
    for r in data.get("included", []):
        app = appraisal_by_id.get(r["Record_ID"], {})
        included_rows.append([
            r.get("Record_ID", ""),
            clean(r.get("Title", ""), 130),
            clean(r.get("Study_Type", ""), 45),
            clean(app.get("Overall_Appraisal", ""), 35),
            clean(app.get("Use_For", ""), 80),
        ])
    add_table(doc, ["ID", "题名", "研究类型", "综合评价", "建议用途"], included_rows, widths=[0.55, 2.35, 1.2, 0.9, 1.5], font_size=7)

    doc.add_heading("8. Pending Confirmation and Limitations / 待确认和局限", level=1)
    add_bullet(doc, f"本轮待人工确认记录数为 {decision_counts.get('待人工确认', 0)}，主要由于题名主题较宽、仅全文提示相关、结局指标自动抽取不充分或报告质量需复核。")
    add_bullet(doc, f"PubMed未匹配但潜在一次性输尿管镜题录数为 {summary.get('potential_single_use_ureteroscope_unmatched_pubmed_count', 0)}，已在Excel中单独列出，后续可决定是否补下载全文。")
    add_bullet(doc, "自动筛选依赖PDF前若干页文本和题录摘要，若核心证据用于正式提交，应逐篇核对全文、表格和不良事件定义。")
    add_bullet(doc, "产品目录文献已按范围排除；若后续需要产品性能横向比较，可另建独立产品证据包。")

    doc.add_heading("9. Conclusion / 结论", level=1)
    doc.add_paragraph(
        f"本轮筛选形成了一个以一次性/可抛弃输尿管镜为核心的证据包：{dataset_counts.get('Similar device clinical dataset', 0)}条相似器械临床数据、{dataset_counts.get('SOTA dataset', 0)}条SOTA数据、{dataset_counts.get('Technical supportive data', 0)}条技术支持数据、{dataset_counts.get('Economic/resource evidence', 0)}条经济/资源证据、{dataset_counts.get('Supportive case evidence', 0)}条病例支持证据，以及{dataset_counts.get('Pending confirmation', 0)}条待确认记录。整体证据可用于支持一次性输尿管镜在URS/RIRS/FURS及上尿路结石诊疗中的安全性、性能、临床价值、资源利用和技术现状论证，但核心结论在正式提交前仍需全文核对。"
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Single-use ureteroscope literature screening report - generated from local screening workbook")
    fr.font.name = "Calibri"
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string("6B7280")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
