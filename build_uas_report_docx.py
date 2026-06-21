from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DATA_JSON = ROOT / "outputs" / "uas_screening" / "uas_screening_data.json"
OUTPUT = ROOT / "outputs" / "uas_screening" / "UAS文献筛选与评价报告_完整版.docx"


def clean(text: str, limit: int | None = None) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    if limit and len(text) > limit:
        return text[: limit - 1].rstrip() + "..."
    return text


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold=False, color="111827", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(clean(str(text)))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="0B2545", size=font_size)
        set_cell_shading(hdr[i], "F2F4F7")
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("111827")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ureteral Access Sheath (UAS) 文献检索、筛选与评价报告")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"基于本地三大数据库及二次搜索文献 | 生成日期：{date.today().isoformat()}")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("4B5563")
    doc.add_paragraph()


def build_report() -> None:
    data = json.loads(DATA_JSON.read_text(encoding="utf-8"))
    summary = data["summary"]
    report = data["report_summary"]
    app_counts = summary.get("appraisal_overall_counts", {})
    dataset_counts = report.get("dataset_counts", {})

    doc = Document()
    setup_styles(doc)
    add_title(doc)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "本报告按照 ureteral access sheath (UAS) 主题口径，对本地Cochrane、Google Scholar、PubMed及二次搜索下载文献进行去重、筛选、数据集分层、适应性评估和安全性/性能证据摘录。"
    )
    add_table(
        doc,
        ["项目", "结果"],
        [
            ["PDF总数", str(summary["pdf_count"])],
            ["去重后筛选记录", str(summary["screened_unique_count"])],
            ["纳入 / 排除 / 待确认", f"{summary['decision_counts'].get('纳入', 0)} / {summary['decision_counts'].get('排除', 0)} / {summary['decision_counts'].get('待人工确认', 0)}"],
            ["相似器械临床数据", str(dataset_counts.get("Similar device clinical dataset", 0))],
            ["SOTA数据", str(dataset_counts.get("SOTA dataset", 0))],
            ["适应性综合评价", "；".join([f"{k}: {v}" for k, v in app_counts.items()])],
        ],
        widths=[2.2, 4.1],
        font_size=9,
    )
    doc.add_paragraph(report["key_message"])

    doc.add_heading("2. Literature Search Scope and Source", level=1)
    add_bullet(doc, "纳入来源：Cochrane、Google Scholar、PubMed，以及文献2次搜索(24-26.6)中的对应数据库文件夹。")
    add_bullet(doc, "不纳入范围：产品文献文件夹。")
    add_bullet(doc, "PubMed TXT题录用于补充题名、作者、期刊、年份、DOI、PMID和摘要；题录与PDF采用保守匹配，低置信匹配不强行采用。")
    add_bullet(doc, "筛选.docx作为流程、纳排记录和适应性评估结构参考；主题标准已从Double-J支架改写为UAS。")

    doc.add_heading("3. Screening Criteria and Process", level=1)
    doc.add_paragraph("纳入条件：题名或摘要显示UAS、ureteral/ureteric access sheath、suction UAS、TFS-UAS、FV-UAS、FANS、SUAS等为研究对象、干预、对照、技术或主要临床问题，并报告安全性、有效性、性能、压力、感染、结石清除、并发症或资源效率等相关结局。")
    doc.add_paragraph("排除条件：无UAS主题；UAS仅在背景、常规操作或参考文献中顺带出现；主要研究对象为输尿镜、PCNL、ESWL、激光、取石篮、支架或其他非UAS器械/策略；或缺少UAS相关安全性/性能结局。")
    add_table(
        doc,
        ["流程步骤", "数量", "说明"],
        [[r["Step"], str(r["Count"]), r["Description"]] for r in data.get("flow_rows", [])],
        widths=[2.0, 0.8, 3.5],
        font_size=8,
    )

    doc.add_heading("4. Dataset Assignment", level=1)
    doc.add_paragraph("为模拟筛选.docx中的SOTA/Similar device证据组织方式，本报告将UAS证据分为以下数据集。")
    add_table(
        doc,
        ["数据集", "数量", "用途"],
        [
            ["Similar device clinical dataset", str(dataset_counts.get("Similar device clinical dataset", 0)), "直接评价UAS或相似UAS的临床安全性/性能，作为核心临床证据。"],
            ["SOTA dataset", str(dataset_counts.get("SOTA dataset", 0)), "系统综述/Meta等二级证据，用于UAS技术现状和总体安全/性能背景。"],
            ["Technical supportive data", str(dataset_counts.get("Technical supportive data", 0)), "体外、工程或模型研究，支持机制或技术性能，临床评价需降权。"],
            ["Supportive case evidence", str(dataset_counts.get("Supportive case evidence", 0)), "病例或特殊场景证据，作补充支持。"],
            ["Pending confirmation", str(dataset_counts.get("Pending confirmation", 0)), "主题权重或证据用途需人工确认。"],
        ],
        widths=[2.0, 0.7, 3.6],
        font_size=8,
    )

    doc.add_heading("5. Suitability and Contribution Appraisal", level=1)
    doc.add_paragraph("适应性评估采用UAS改写版D/A/P/R/T/O/S/C框架，分别评价器械相关性、用途一致性、患者适用性、报告质量、证据类型、结局相关性、统计支持和临床意义。")
    add_table(
        doc,
        ["综合评价", "数量", "解释"],
        [
            ["高/核心证据", str(app_counts.get("高/核心证据", 0)), "器械、用途、人群、报告和结局均与UAS临床评价高度匹配。"],
            ["中等/可作为支持证据", str(app_counts.get("中等/可作为支持证据", 0)), "与UAS相关，但研究类型、患者人群、报告质量或统计/临床意义存在限制。"],
            ["低/仅支持性", str(app_counts.get("低/仅支持性", 0)), "非临床、病例或临床意义有限，需降权使用。"],
            ["待确认", str(app_counts.get("待确认", 0)), "题名/摘要提示相关，但主题权重或证据用途需人工复核。"],
        ],
        widths=[1.9, 0.7, 3.7],
        font_size=8,
    )

    doc.add_heading("6. Safety and Performance Evidence Summary", level=1)
    doc.add_paragraph("自动数据摘录显示，UAS相关文献主要围绕结石清除率、术后并发症、感染/发热/SIRS、手术时间、住院时间、出血或血红蛋白下降、肾内压力/灌注/负压吸引效率等维度报告安全性和性能。详细逐篇摘录见Excel工作簿“安全性能数据提取”sheet。")
    top_evidence = [r for r in data.get("performance_safety_extraction", []) if r.get("Dataset_Assignment") == "Similar device clinical dataset"][:10]
    add_table(
        doc,
        ["ID", "题名", "样本量", "关键性能/安全性摘录"],
        [
            [
                r["Record_ID"],
                clean(r["Title"], 120),
                r.get("Sample_Size_N", ""),
                clean(" ".join([r.get("Performance_Outcomes", ""), r.get("Safety_Outcomes", "")]), 260),
            ]
            for r in top_evidence
        ],
        widths=[0.6, 2.3, 0.7, 2.7],
        font_size=7,
    )

    doc.add_heading("7. Included Evidence List", level=1)
    included_rows = []
    appraisal_by_id = {r["Record_ID"]: r for r in data.get("appraisal", [])}
    for r in data.get("included", []):
        app = appraisal_by_id.get(r["Record_ID"], {})
        included_rows.append([
            r["Record_ID"],
            clean(r["Title"], 120),
            r["Study_Type"],
            app.get("Overall_Appraisal", ""),
            app.get("Use_For", ""),
        ])
    add_table(doc, ["ID", "题名", "研究类型", "综合评价", "建议用途"], included_rows, widths=[0.55, 2.4, 1.2, 0.9, 1.25], font_size=7)

    doc.add_heading("8. Pending Confirmation and Limitations", level=1)
    pending_rows = [[r["Record_ID"], clean(r["Title"], 140), r["Inclusion_Class"], clean(r["Notes"], 120)] for r in data.get("pending", [])]
    add_table(doc, ["ID", "题名", "类型", "待确认原因"], pending_rows, widths=[0.55, 2.8, 1.4, 1.55], font_size=7)
    add_bullet(doc, "自动摘录依赖PDF前12页可读文本和PubMed题录，正式提交前建议对高/核心证据逐篇全文核对。")
    add_bullet(doc, "部分Google Scholar/Cochrane PDF缺少题录元数据，题名、年份和DOI可能需人工补全。")
    add_bullet(doc, "PubMed中有潜在UAS题录未匹配到本地PDF，已在Excel中单独列出，后续可决定是否补下载全文。")
    add_bullet(doc, "非临床、病例和宽泛指南类证据可作为支持或背景，不建议作为核心临床安全性/性能证据。")

    doc.add_heading("9. Conclusion", level=1)
    doc.add_paragraph(
        "本轮筛选形成了一个以UAS为核心的证据包：13条相似器械临床数据可作为主要临床证据，4条SOTA数据可支撑技术现状和总体安全/性能背景，另有技术支持、病例支持和待确认记录。整体证据提示，吸引、可弯曲或导航UAS在RIRS/FURS处理肾及上尿路结石时主要贡献集中在结石清除、压力/灌注管理、并发症和感染风险控制、手术效率与住院资源等方面。"
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("UAS literature screening report - generated from local screening workbook")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string("6B7280")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
